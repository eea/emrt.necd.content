try:
    from cStringIO import StringIO
except ImportError:
    from StringIO import StringIO
import datetime
import json
import re
from docx import Document
from docx.shared import Pt
from docx.enum.style import WD_STYLE_TYPE
from AccessControl import getSecurityManager
from Acquisition import aq_inner
from emrt.necd.content.roles.localrolesubscriber import grant_local_roles
from plone import api
from plone.app.contentlisting.interfaces import IContentListing
from plone.dexterity.browser import add
from plone.dexterity.browser import edit
from plone.dexterity.browser.view import DefaultView
from plone.app.discussion.interfaces import IConversation
from plone.dexterity.content import Container
from plone.directives import form
from plone.memoize import instance
from plone.namedfile.interfaces import IImageScaleTraversable
from plone.z3cform.interfaces import IWrappedForm
from Products.CMFCore.utils import getToolByName
from Products.CMFEditions import CMFEditionsMessageFactory as _CMFE
from Products.CMFPlone.utils import safe_unicode
from Products.Five import BrowserView
from Products.statusmessages.interfaces import IStatusMessage
from time import time
from z3c.form import button
from z3c.form import field
from z3c.form import interfaces
from z3c.form import validator
from z3c.form.browser.checkbox import CheckBoxFieldWidget
from z3c.form.form import Form
from z3c.form.interfaces import ActionExecutionError
import zope.schema as schema
from zope.browsermenu.menu import getMenu
from zope.browserpage.viewpagetemplatefile import (
    ViewPageTemplateFile as Z3ViewPageTemplateFile
)
from zope.component import getUtility
from zope.i18n import translate
from zope.interface import alsoProvides
from zope.interface import implementer
from zope.interface import Invalid
from emrt.necd.content import MessageFactory as _
from eea.cache import cache
from .comment import IComment
from .commentanswer import ICommentAnswer
from .nfr_code_matching import get_category_ldap_from_nfr_code
from .nfr_code_matching import get_category_value_from_nfr_code
from emrt.necd.content.subscriptions.interfaces import (
    INotificationUnsubscriptions
)
from emrt.necd.content.utilities import ms_user
from emrt.necd.content.constants import LDAP_SECTOREXP
from emrt.necd.content.constants import ROLE_SE
from emrt.necd.content.constants import ROLE_CP
from emrt.necd.content.constants import ROLE_LR
from emrt.necd.content.constants import P_OBS_REDRAFT_REASON_VIEW
from emrt.necd.content.utils import activity_data_validator
from emrt.necd.content.utils import get_vocabulary_value
from emrt.necd.content.utils import hidden
from emrt.necd.content.utilities.interfaces import IFollowUpPermission
from plone.supermodel import model

from emrt.necd.content.vocabularies import get_registry_interface_field_data
from emrt.necd.content.vocabularies import INECDVocabularies


YEAR_DESCRIPTION_PROJECTION = u"Inventory year is the year or a list of " \
                              u"years (e.g. '2050', '2020, 2025, 2030') when" \
                              u" the emissions had occured for which an issue" \
                              u" was observed in the review. The allowed " \
                              u"values are: 2020, 2025, 2030, 2040 or 2050."

YEAR_DESCRIPTION_INVENTORY = u"Inventory year is the year, a range or a list " \
                             u"of years (e.g. '2012', '2009-2012', " \
                             u"'2009, 2012, 2013') when the emissions had " \
                             u"occured for which an issue was observed in the" \
                             u" review."


# Cache helper methods
def _user_name(fun, self, userid):
    return (userid, time() // 86400)

def _is_projection(context):
    return context.type == 'projection'


def check_parameter(value):
    if len(value) == 0:
        raise Invalid(u'You need to select at least one parameter')

    return True


def check_nfr_code(value):
    """ Check if the user is in one of the group of users
        allowed to add this category NFR Code observations
    """
    category = get_category_ldap_from_nfr_code(value)
    user = api.user.get_current()
    groups = user.getGroups()
    valid = False
    for group in groups:
        if group.startswith('{}-{}-'.format(LDAP_SECTOREXP, category)):
            valid = True

    if not valid:
        raise Invalid(
            u'You are not allowed to add observations for this sector category'
        )

    return True


def check_country(value):
    user = api.user.get_current()
    groups = user.getGroups()
    valid = False
    for group in groups:
        if group.startswith('{}-'.format(LDAP_SECTOREXP)) and \
                group.endswith('-%s' % value):
            valid = True

    if not valid:
        raise Invalid(
            u'You are not allowed to add observations for this country'
        )

    return True


def inventory_year(value):
    """
    Inventory year can be a given year (2014), a range of years (2012-2014)
    or a list of the years (2012, 2014, 2016)
    """

    def split_on_sep(val, sep):
        for s in sep:
            if s in val:
                return tuple(val.split(s))
        return (val, )

    def validate(value):
        normalized_value = (val.strip() for val in split_on_sep(value, '-,;'))
        return False not in (int(val) > 0 for val in normalized_value)

    def check_valid(value):
        try:
            return validate(value)
        except ValueError:
            return False

    if not check_valid(value):
        raise Invalid(u'Inventory year format is not correct. ')

    return True


def default_year():
    return datetime.datetime.now().year


# Interface class; used to define content-type schema.
class IObservation(model.Schema, IImageScaleTraversable):
    """
    New review observation
    """

    text = schema.Text(
        title=u'Short description by sector expert',
        required=True,
        description=(
            u"Describe the issue identified. Keep it short, you cannot "
            u"change this description once you have sent it to LR. MS can "
            u"only see the question once it has been approved and sent by "
            u"the LR. The question to the MS should be asked in the Q&A tab, "
            u"not here."
        )
    )

    country = schema.Choice(
        title=u"Country",
        vocabulary='emrt.necd.content.eea_member_states',
        constraint=check_country,
        required=True,
    )

    nfr_code = schema.Choice(
        title=u"NFR category codes",
        vocabulary='emrt.necd.content.nfr_code',
        constraint=check_nfr_code,
        required=True,
    )

    year = schema.TextLine(
        title=u'Inventory year',
        required=True,
    )

    form.widget(pollutants=CheckBoxFieldWidget)
    pollutants = schema.List(
        title=u"Pollutants",
        value_type=schema.Choice(
            vocabulary='emrt.necd.content.pollutants',
        ),
        required=True,
    )

    review_year = schema.Int(
        title=u'Review year',
        description=u'Review year is the year in which the inventory was ' \
                    u'submitted and the review was carried out',
        defaultFactory=default_year,
        required=True,
    )

    fuel = schema.Choice(
        title=u"Fuel",
        vocabulary='emrt.necd.content.fuel',
        required=False,
    )

    activity_data_type = schema.Choice(
        title=u"Activity Data Type",
        vocabulary='emrt.necd.content.activity_data_type',
        required=False,
    )

    form.widget(activity_data=CheckBoxFieldWidget)
    activity_data = schema.List(
        title=u"Activity Data",
        value_type=schema.Choice(
            vocabulary='emrt.necd.content.activity_data',
        ),
        required=False,
    )

    ms_key_category = schema.Bool(
        title=u"MS key category",
    )

    form.widget(parameter=CheckBoxFieldWidget)
    parameter = schema.List(
        title=u"Parameter",
        value_type=schema.Choice(
            vocabulary='emrt.necd.content.parameter',
        ),
        constraint=check_parameter,
        required=True,
    )

    form.widget(highlight=CheckBoxFieldWidget)
    highlight = schema.List(
        title=u"Description flags",
        description=(
            u"Description flags highlight important information "
            u"that is closely related to the item."
        ),
        value_type=schema.Choice(
            vocabulary='emrt.necd.content.highlight',
        ),
        required=False,
        default=[],
    )

    form.write_permission(closing_comments='cmf.ManagePortal')
    closing_comments = schema.Text(
        title=u'Finish request comments',
        required=False,
    )

    form.write_permission(closing_deny_comments='cmf.ManagePortal')
    closing_deny_comments = schema.Text(
        title=u'Finish deny comments',
        required=False,
    )


@form.validator(field=IObservation['parameter'])
def check_parameter(value):
    if len(value) == 0:
        raise Invalid(u'You need to select at least one parameter')


@form.validator(field=IObservation['pollutants'])
def check_pollutants(value):
    if len(value) == 0:
        raise Invalid(u'You need to select at least one pollutant')


@form.validator(field=IObservation['country'])
def check_country(value):
    user = api.user.get_current()
    groups = user.getGroups()
    valid = False
    for group in groups:
        if group.startswith('{}-'.format(LDAP_SECTOREXP)) and \
                group.endswith('-%s' % value):
            valid = True

    if not valid:
        raise Invalid(
            u'You are not allowed to add observations for this country'
        )


@form.validator(field=IObservation['year'])
def inventory_year(value):
    """
    Inventory year can be a given year (2014), a range of years (2012-2014)
    or a list of the years (2012, 2014, 2016).
    In the case of a 'Projection' ReviewFolder the values must be:
    2020, 2025, 2030, 2040 or 2050.
    """
    def split_on_sep(val, sep):
        for s in sep:
            if s in val:
                return tuple(val.split(s))
        return (val, )

    def validate(value):
        normalized_value = (val.strip() for val in split_on_sep(value, '-,;'))
        return False not in (int(val) > 0 for val in normalized_value)

    def check_valid(value):
        try:
            return validate(value)
        except ValueError:
            return False

    if not check_valid(value):
        raise Invalid(u'Inventory year format is not correct. ')


class NfrCodeContextValidator(validator.SimpleFieldValidator):
    def validate(self, value, force=False):
        """ Check if the user is in one of the group of users
            allowed to add this category NFR Code observations
        """
        category = get_category_ldap_from_nfr_code(value, self.context)
        user = api.user.get_current()
        groups = user.getGroups()
        valid = False
        for group in groups:
            if group.startswith('{}-{}-'.format(LDAP_SECTOREXP, category)):
                valid = True

        if not valid:
            raise Invalid(
                u'You are not allowed to add observations for this sector category'
            )


validator.WidgetValidatorDiscriminators(
    NfrCodeContextValidator,
    field=IObservation['nfr_code']
)


class InventoryYearContextValidator(validator.SimpleFieldValidator):
    def validate(self, value, force=False):

        if _is_projection(self.context):
            allowed_years = ['2020', '2025', '2030', '2040', '2050']

            if value is None :
                # Assume empty string = no input
                return

            value = [val.strip() for val in value.split(',')]

            if not set(value).issubset(allowed_years):
                raise Invalid(_(u'The entered value is not correct.'))


validator.WidgetValidatorDiscriminators(
    InventoryYearContextValidator,
    field=IObservation['year']
)


def set_title_to_observation(object, event):
    sector = safe_unicode(object.ghg_source_category_value())
    pollutants = safe_unicode(object.pollutants_value())
    inventory_year = safe_unicode(str(object.year))
    parameter = safe_unicode(object.parameter_value())
    object.title = u' '.join([sector, pollutants, inventory_year, parameter])
    grant_local_roles(object)


@implementer(IObservation)
class Observation(Container):

    def get_values(self):
        """
        Memoized version of values, to speed-up
        """
        return self.values()

    def get_values_cat(self, portal_type=None):
        if portal_type is not None:
            return self.listFolderContents(
                contentFilter={'portal_type': portal_type}
            )
        else:
            return self.listFolderContents()


    def get_nfr_code(self):
        """ stupid method to avoid name-clashes with the existing
        vocabularies when cataloging """
        return self.nfr_code

    def get_ghg_source_sectors(self):
        """ stupid method to avoid name-clashes with the existing
        vocabularies when cataloging """
        return self.ghg_source_sectors_value()

    def get_highlight(self):
        """ stupid method to avoid name-clashes with the existing
        vocabularies when cataloging """
        return self.highlight

    def country_value(self):
        return get_vocabulary_value(
            self, 'emrt.necd.content.eea_member_states', self.country
        )

    def nfr_code_value(self):
        # Same as ghg_source_sectors_value as the object reindex
        # might be triggered by a user that doesn't have access
        # to the current nfr code, resulting in an empty value
        # getting indexed.
        return self.ghg_source_sectors_value()

    def ghg_source_category_value(self):
        # Get the value of the sector to be used on the LDAP mapping
        return get_category_ldap_from_nfr_code(self.nfr_code, self.aq_parent)

    def ghg_source_sectors_value(self):
        # Get the value of the sector to be used
        # on the Observation Metadata screen
        return get_category_value_from_nfr_code(self.nfr_code, self.aq_parent)

    def parameter_value(self):
        parameters = [
            get_vocabulary_value(self, 'emrt.necd.content.parameter', p)
            for p in self.parameter
            ]
        return u', '.join(filter(None, parameters))

    def pollutants_value(self):
        pollutants = [
            get_vocabulary_value(self, 'emrt.necd.content.pollutants', p)
            for p in self.pollutants
            ]
        return u', '.join(filter(None, pollutants))

    def activity_data_value(self):
        activities = [
            get_vocabulary_value(self, 'emrt.necd.content.activity_data', ad)
            for ad in self.activity_data
            ]
        return activities

    def highlight_value(self):
        if self.highlight:
            highlight = [
                get_vocabulary_value(self,'emrt.necd.content.highlight',h)
                for h in self.highlight
                ]
            return u', '.join(highlight)
        return u''

    def finish_reason_value(self):
        return get_vocabulary_value(self,
            'emrt.necd.content.finishobservationreasons',
            self.closing_reason
        )

    def finish_deny_reason_value(self):
        return get_vocabulary_value(
            self, 'emrt.necd.content.finishobservationdenyreasons',
            self.closing_deny_reason
        )

    def get_status(self):
        return api.content.get_state(self)

    def can_draft_conclusions(self):
        questions = self.get_values_cat('Question')
        if len(questions) > 0:
            q = questions[0]
            return q.get_state_api() in [
                'draft',
                'drafted',
                'recalled-lr',
                'closed',
            ]
        else:
            return True

    def can_close(self):
        if self.get_status() in ['pending']:
            questions = self.get_values_cat('Question')
            if len(questions) > 0:
                for q in questions:
                    if q.get_state_api() not in ['closed']:
                        return False
                return True

        return False

    def wf_location(self):
        status = self.get_status()
        if status in ('draft', 'conclusions', 'conclusions-lr-denied'):
            return 'Sector Expert'
        elif status == 'closed':
            return 'Lead reviewer'
        elif status == 'conclusion-discussion':
            return 'Counterpart'
        elif status  == 'close-requested':
            return 'Lead reviewer'
        else:
            questions = self.get_values_cat('Question')
            if questions:
                question = questions[0]
                state = question.get_state_api()
                if state in ['draft', 'closed']:
                    return 'Sector Expert'
                elif state in ['counterpart-comments']:
                    return 'Counterparts'
                elif state in ['drafted', 'recalled-lr']:
                    return 'Lead reviewer'
                elif state in [
                        'pending',
                        'answered',
                        'pending-answer-drafting',
                        'recalled-msa']:
                    return 'Member state coordinator'
                elif state in ['expert-comments']:
                    return 'Member state experts'
            else:
                return "Sector Expert"

    def wf_status(self):
        if self.get_status() in ['draft']:
            return ['Observation created', "observationIcon"]
        elif self.get_status() in ['closed']:
            return ['Observation finished', "observationIcon"]
        elif self.get_status() in ['close-requested']:
            return ['Observation finish requested', "observationIcon"]
        elif self.get_status() in ['conclusions', 'conclusions-lr-denied']:
            return ["Conclusion ongoing", "conclusionIcon"]
        elif self.get_status() in ['conclusion-discussion']:
            return ["Counterparts comments requested", "conclusionIcon"]
        else:
            questions = self.get_values_cat('Question')
            if questions:
                question = questions[-1]
                state = question.get_state_api()
                if state in ['draft']:
                    return ["Question drafted", "questionIcon"]
                elif state in ['counterpart-comments']:
                    return ["Counterpart's comments requested", "questionIcon"]
                elif state in ['answered']:
                    return ['Pending question', "questionIcon"]
                elif state in [
                        'pending',
                        'pending-answer-drafting',
                        'recalled-msa']:
                    return ['Open question', "questionIcon"]
                elif state in ['drafted', 'recalled-lr']:
                    return ['Draft question', "questionIcon"]
                elif state in ['expert-comments']:
                    return ['MS expert comments requested', 'questionIcon']
                elif state in ['closed']:
                    return ['Closed question', "questionIcon"]
            else:
                return ['Observation created', "observationIcon"]

        return ['Unknown', 'observationIcon']

    def observation_status(self):
        status = self.observation_question_status()
        my_status = self.get_status()

        if status in ['draft',
                      'counterpart-comments',
                      'observation-draft']:
            return 'SE'
        elif status in ['drafted', 'recalled-lr']:
            return 'LR'
        elif status in ['pending',
                        'pending-answer-drafting',
                        'expert-comments']:
            return 'MSC'
        elif status in ['answered']:
            return 'answered'
        elif status in ['conclusions',
                        'conclusions-lr-denied',
                        'conclusion-discussion']:
            return 'conclusions'
        elif status in ['close-requested']:
            return 'close-requested'
        elif status == 'closed':
            conclusion = self.get_conclusion()
            conclusion_reason = conclusion and conclusion.closing_reason or ''
            if (conclusion_reason == 'no-conclusion-yet'):
                return "SE"
            elif not my_status.endswith('closed'):
                return "answered"
            else:
                return "finalised"
        else:
            return status


    def observation_questions_workflow(self):
        questions = self.get_values_cat('Question')
        if not questions:
            return tuple()

        # there is always only one question.
        question = questions[0]

        items = question.values()

        comments = [i for i in items if i.portal_type == 'Comment']
        answers = [i for i in items if i.portal_type == 'CommentAnswer']

        len_comments = len(comments)
        obs_status = self.observation_status()

        return tuple(['Answered'] * (len_comments - 1) + [obs_status])


    def overview_status(self):
        status = self.get_status()
        closed_val = 'closed ({reason})'
        if status == 'closed':
            conclusion = self.get_conclusion()
            if conclusion:
                return closed_val.format(reason=conclusion.reason_value())
        else:
            return 'open'

    def is_secretariat(self):
        user = api.user.get_current()
        return 'Manager' in user.getRoles()

    @cache(_user_name)
    def _author_name(self, userid):
        if userid:
            user = api.user.get(username=userid)
            if user:
                return user.getProperty('fullname', userid)

        return userid

    def get_author_name(self, userid=None):
        if not userid:
            userid = self.Creator()

        return self._author_name(userid)

    def myHistory(self):
        observation_history = self.workflow_history.get(
            'esd-review-workflow', [])
        observation_wf = []
        question_wf = []
        for item in observation_history:
            item['role'] = item['actor']
            item['object'] = 'observationIcon'
            item['author'] = self.get_author_name(item['actor'])
            if item['review_state'] == 'draft':
                item['state'] = 'Draft observation'
                item['role'] = "Sector Expert"
                observation_wf.append(item)
            elif item['review_state'] == 'pending' and item['action'] == "approve":
                item['state'] = 'Pending'
                # Do not add
            elif item['review_state'] == 'pending' and item['action'] == "reopen":
                item['state'] = 'Observation reopened'
                item['role'] = "Sector Expert"
                observation_wf.append(item)
            elif item['review_state'] == 'pending':
                item['state'] = 'Pending'
                # Do not add
            elif item['review_state'] == 'closed':
                item['state'] = 'Closed observation'
                item['role'] = "Lead Reviewer"
                observation_wf.append(item)
            elif item['review_state'] == 'close-requested':
                item['state'] = 'Finalisation requested'
                item['role'] = "Sector Expert"
                observation_wf.append(item)
            elif item['review_state'] in ('conclusions', 'conclusions-lr-denied') and item['action'] == "deny-finishing-observation":
                item['state'] = 'Finalisation denied'
                item['role'] = "Lead reviewer"
                observation_wf.append(item)
            elif item['review_state'] == 'conclusions-lr-denied' and item['action'] == "recall-lr":
                item['state'] = 'Recalled by LR'
                item['role'] = "Lead reviewer"
                observation_wf.append(item)
            elif item['review_state'] == 'closed' and item['action'] == "recall-lr":
                item['state'] = 'Recalled by LR'
                item['role'] = "Lead reviewer"
                observation_wf.append(item)
            elif item['review_state'] == 'conclusion-discussion':
                item['state'] = 'Conclusion comments requested'
                item['role'] = "Sector Expert"
                item['object'] = 'conclusionIcon'
                observation_wf.append(item)
            elif item['review_state'] == 'conclusions' and item['action'] == "finish-comments":
                item['state'] = 'Conclusion comments closed'
                item['role'] = "Sector Expert"
                item['object'] = 'conclusionIcon'
                observation_wf.append(item)
            elif item['review_state'] == 'conclusions' and item['action'] == "draft-conclusions":
                item['state'] = 'Conclusion drafting'
                item['role'] = "Sector Expert"
                item['object'] = 'conclusionIcon'
                observation_wf.append(item)
            else:
                item['state'] = '*' + item['review_state'] + '*'
                observation_wf.append(item)

        history = list(observation_wf)
        questions = self.get_values_cat()

        if questions:
            question = questions[0]
            question_history = question.workflow_history.get(
                'esd-question-review-workflow', [])
            for item in question_history:
                item['role'] = item['actor']
                item['object'] = 'questionIcon'
                item['author'] = self.get_author_name(item['actor'])
                if item['review_state'] == 'draft' and item['action'] == "reopen":
                    item['state'] = 'Draft question'
                    item['role'] = "Sector Expert"
                    question_wf.append(item)
                elif item['review_state'] == 'counterpart-comments':
                    item['state'] = 'Requested counterparts comments'
                    item['role'] = "Sector Expert"
                    question_wf.append(item)
                elif item['review_state'] == 'draft' and item['action'] =='send-comments':
                    item['state'] = 'Counterparts comments closed'
                    item['role'] = "Sector Expert"
                    question_wf.append(item)
                elif item['review_state'] == 'drafted':
                    item['state'] = 'Sent to LR'
                    item['role'] = "Sector Expert"
                    question_wf.append(item)
                elif item['review_state'] == 'draft' and item['action'] == 'recall-sre':
                    item['state'] = 'Question recalled'
                    item['role'] = "Sector Expert"
                    question_wf.append(item)
                elif item['review_state'] == 'draft' and item['action'] == 'redraft':
                    item['state'] = 'Question redrafted'
                    item['role'] = "Sector Expert"
                    question_wf.append(item)
                elif item['review_state'] == 'draft':
                    # Do not add
                    pass
                elif item['review_state'] == 'pending' and item['action'] == 'approve-question':
                    item['state'] = 'Question approved and sent to MS coordinator'
                    item['role'] = "Lead reviewer"
                    question_wf.append(item)
                elif item['review_state'] == 'recalled-lr':
                    item['state'] = 'Question recalled'
                    item['role'] = "Lead reviewer"
                elif item['review_state'] == 'answered':
                    item['state'] = 'Answer sent'
                    item['role'] = "Member state coordinator"
                    question_wf.append(item)
                elif item['review_state'] == 'expert-comments':
                    item['state'] = 'MS expert comments requested'
                    item['role'] = "Member state coordinator"
                    question_wf.append(item)
                elif item['review_state'] == 'pending-answer-drafting':
                    item['state'] = 'Member state expert comments closed'
                    item['role'] = "Member state coordinator"
                    question_wf.append(item)
                elif item['review_state'] == 'recalled-msa':
                    item['state'] = 'Answer recalled'
                    item['role'] = "Member state coordinator"
                    question_wf.append(item)
                elif item['action'] == 'validate-answer-msa' and item['action'] == 'validate-answer-msa':
                    item['state'] = 'Answer acknowledged'
                    item['role'] = "Sector Expert"
                    question_wf.append(item)
                elif item['review_state'] == 'draft' and item['action'] == "reopen":
                    item['state'] = 'Reopened'
                    # Do not add
                elif item['review_state'] == 'closed':
                    item['state'] = 'Closed'
                    # Do not add
                else:
                    item['state'] = '*' + item['review_state'] + '*'
                    item['role'] = item['actor']
                    question_wf.append(item)

            history = list(observation_wf) + list(question_wf)

        history.sort(key=lambda x: x["time"], reverse=False)
        return history

    def can_edit(self):
        sm = getSecurityManager()
        return sm.checkPermission('Modify portal content', self)

    def get_question(self):
        questions = self.get_values_cat('Question')

        if questions:
            question = questions[-1]
            return question

    def observation_question_status(self):
        questions = self.get_values_cat('Question')
        if self.get_status() != 'pending':
            if self.get_status() in ['conclusions', 'conclusions-lr-denied']:
                if questions:
                    question = questions[-1]
                    question_state = api.content.get_state(question)
                    if question_state != 'closed':
                        return question_state
            return self.get_status()
        else:
            if questions:
                question = questions[-1]
                state = api.content.get_state(question)
                return state
            else:
                return "observation-draft"

    def observation_css_class(self):
        if self.highlight:
            if self.get_status() == "closed":
                con = self.get_conclusion()
                if con:
                    if con.closing_reason == "technical-correction":
                        return 'technicalCorrectionBackground'

            elif 'ptc' in self.highlight:
                return 'ptcBackground'

    def observation_is_potential_significant_issue(self):
        if self.highlight:
            return 'psi' in self.highlight
        return False

    def observation_is_potential_technical_correction(self):
        if self.highlight:
            return 'ptc' in self.highlight
        return False

    def observation_is_technical_correction(self):
        if self.get_status() == "closed":
            con = self.get_conclusion()
            if con:
                return con.closing_reason == "technical-correction"
        return False

    def get_conclusion(self):
        conclusions = self.get_values_cat('Conclusions')
        mtool = api.portal.get_tool('portal_membership')
        if conclusions and mtool.checkPermission('View', conclusions[0]):
            return conclusions[0]
        return None

    def last_question_reply_number(self):
        questions = self.get_values_cat('Question')
        replynum = 0
        if questions:
            comments = [c for c in questions[-1].values() if c.portal_type == "Comment"]
            if comments:
                last = comments[-1]
                disc = IConversation(last)
                return disc.total_comments

        return replynum

    def last_answer_reply_number(self):
        questions = self.get_values_cat('Question')
        replynum = 0
        if questions:
            comments = [c for c in questions[-1].values() if c.portal_type == "CommentAnswer"]
            if comments:
                last = comments[-1]
                disc = IConversation(last)
                return disc.total_comments

        return replynum

    def reply_comments_by_mse(self):
        questions = self.get_values_cat('Question')
        user = api.user.get_current().id
        if questions:
            comments = [c for c in questions[-1].values() if c.portal_type == "CommentAnswer"]
            if comments:
                last = comments[-1]
                disc = IConversation(last)
                return user in IConversation(last).commentators

        return False

    def observation_already_replied(self):

        questions = self.get_values_cat('Question')
        if questions:
            question = questions[0]
            winfo = question.workflow_history
            state = self.get_status()
            for witem in winfo.get('esd-question-review-workflow', []):
                if witem.get('review_state', '') == 'answered':
                    return True

        return False

    def can_add_followup(self):
        status = self.get_status()
        return status in ['conclusions', 'conclusions-lr-denied']


# View class
# The view will automatically use a similarly named template in
# templates called observationview.pt .
# Template filenames should be all lower case.
# The view will render when you request a content object with this
# interface with "/@@view" appended unless specified otherwise
# This will make this view the default view for your content-type


class EditForm(edit.DefaultEditForm):
    def updateWidgets(self):
        super(EditForm, self).updateWidgets()

        self.widgets['highlight'].template = Z3ViewPageTemplateFile(
            'templates/widget_highlight.pt'
        )
        self.widgets['pollutants'].template = Z3ViewPageTemplateFile(
            'templates/widget_pollutants.pt'
        )

class AddForm(add.DefaultAddForm):
    label = 'Observation'
    description = ' '

    def updateWidgets(self):
        super(AddForm, self).updateWidgets()
        self.fields['IDublinCore.title'].field.required = False

        w_activity_data = self.widgets['activity_data']

        if _is_projection(self.context):
            self.fields['year'].field.description = YEAR_DESCRIPTION_PROJECTION
            self.widgets['fuel'].mode = interfaces.HIDDEN_MODE
            self.widgets['activity_data_type'].template = \
                Z3ViewPageTemplateFile('templates/widget_activity_type.pt')

            w_activity_data.template = Z3ViewPageTemplateFile(
                'templates/widget_activity.pt'
            )
            w_activity_data.activity_data_registry = json.dumps(
                get_registry_interface_field_data(
                    INECDVocabularies,
                    'activity_data'
                )
            )
        else:
            self.fields['year'].field.description = YEAR_DESCRIPTION_INVENTORY
            w_activity_data.mode = interfaces.HIDDEN_MODE
            self.widgets['activity_data_type'].mode = interfaces.HIDDEN_MODE
            self.widgets['pollutants'].template = Z3ViewPageTemplateFile(
                'templates/widget_pollutants.pt'
            )

        self.widgets['IDublinCore.title'].mode = interfaces.HIDDEN_MODE
        self.widgets['IDublinCore.description'].mode = interfaces.HIDDEN_MODE
        self.widgets['text'].rows = 15
        self.widgets['highlight'].template = Z3ViewPageTemplateFile(
            'templates/widget_highlight.pt'
        )
        self.groups = [
            g for g in self.groups if
            g.label == 'label_schema_default'
        ]

    def updateActions(self):
        super(AddForm, self).updateActions()
        self.actions['save'].title = u'Save Observation'
        self.actions['save'].addClass('defaultWFButton')
        self.actions['cancel'].title = u'Delete Observation'

        for k in self.actions.keys():
            self.actions[k].addClass('standardButton')

    def create(self, data={}):
        if _is_projection(self.context):

            activity_data = data['activity_data']
            activity_data_type = data['activity_data_type']

            activity_data_validator(self.context, activity_data_type,
                                    activity_data)


        return super(AddForm, self).create(data)


class AddView(add.DefaultAddView):
    form = AddForm


class ObservationMixin(DefaultView):

    @property
    def user_roles(self):
        user = api.user.get_current()
        return api.user.get_roles(
            username=user.getId(), obj=self.context
        )

    def wf_info(self):
        context = aq_inner(self.context)
        wf = getToolByName(context, 'portal_workflow')
        comments = wf.getInfoFor(self.context,
            'comments', wf_id='esd-review-workflow')
        actor = wf.getInfoFor(self.context,
            'actor', wf_id='esd-review-workflow')
        tim = wf.getInfoFor(self.context,
            'time', wf_id='esd-review-workflow')
        return {'comments': comments, 'actor': actor, 'time': tim}

    def isManager(self):
        sm = getSecurityManager()
        context = aq_inner(self.context)
        return sm.checkPermission('Manage portal', context)

    def get_user_name(self, userid):
        # Check users roles
        country = self.context.country_value()
        sector = self.context.ghg_source_sectors_value()
        return ' - '.join([country, sector])

    def get_menu_actions(self):
        context = aq_inner(self.context)
        menu_items = getMenu(
            'plone_contentmenu_workflow',
            context,
            self.request
        )
        return [mitem for mitem in menu_items if not hidden(mitem)]

    def get_questions(self):
        return IContentListing(self.context.get_values_cat('Question'))

    def can_delete_observation(self):
        is_draft = self.context.get_status() in ['pending', 'draft']
        questions = len(self.context.get_values_cat('Question'))
        #If observation has conclusion cannot be deleted (Ticket #26992)
        conclusions = len(self.context.get_values_cat('Conclusions'))
        return is_draft and not questions and not conclusions

    def can_add_question(self):
        sm = getSecurityManager()
        questions = len(self.context.get_values_cat('Question'))
        return sm.checkPermission('emrt.necd.content: Add Question', self) and not questions

    def can_edit(self):
        sm = getSecurityManager()
        #If observation has conclusion cannot be edited (Ticket #26992)
        conclusions = len(self.context.get_values_cat('Conclusions'))
        return sm.checkPermission('Modify portal content', self.context) and not conclusions

    def get_conclusion(self):
        sm = getSecurityManager()
        conclusions = self.context.get_values_cat('Conclusions')
        if conclusions and sm.checkPermission('View', conclusions[0]):
            return conclusions[0]

        return None

    def existing_conclusion(self):
        conclusion = self.get_conclusion()
        return conclusion and True or False

    def can_add_conclusion(self):
        sm = getSecurityManager()
        question_state = api.content.get_state(self.question())

        return sm.checkPermission(
            'emrt.necd.content: Add Conclusions', self.context
        ) and question_state in ['draft', 'drafted', 'pending', 'closed']

    def show_description(self):
        questions = self.get_questions()
        sm = getSecurityManager()
        if questions:
            question = questions[-1]
            return sm.checkPermission('View', question.getObject())
        else:
            return ms_user.hide_from_ms(self.context)

    def show_internal_notes(self):
        return ms_user.hide_from_ms(self.context)

    def can_view_redraft_reason(self):
        sm = getSecurityManager()
        return sm.checkPermission(P_OBS_REDRAFT_REASON_VIEW, self.context)

    def add_question_form(self):
        form_instance = AddQuestionForm(self.context, self.request)
        alsoProvides(form_instance, IWrappedForm)
        return form_instance()

    def has_local_notifications_settings(self):
        user = api.user.get_current()
        adapted = INotificationUnsubscriptions(self.context)
        data = adapted.get_user_data(user.getId())

        return data and True or False

    # Question view
    def question(self):
        questions = self.get_questions()
        if questions:
            return questions[0].getObject()

    def get_chat(self):
        sm = getSecurityManager()
        question = self.question()
        if question:
            values = [v for v in question.values() if sm.checkPermission('View', v)]
            return values

    def actions(self):
        context = aq_inner(self.context)
        question = self.question()
        observation_menu_items = getMenu(
            'plone_contentmenu_workflow',
            context,
            self.request
        )
        menu_items = observation_menu_items
        if question:
            question_menu_items = getMenu(
                'plone_contentmenu_workflow',
                question,
                self.request
            )
            #remove add-followup-question action if the permission check is False
            if not self.can_add_comment():
                question_menu_items = [item for item in question_menu_items
                                       if not item['action'].endswith('add-followup-question')]

            menu_items = question_menu_items + observation_menu_items
        return [mitem for mitem in menu_items if not hidden(mitem)]

    def get_user_name(self, userid, question=None):
        # check users
        if question is not None:
            country = self.context.country_value()
            sector = self.context.ghg_source_sectors_value()
            if question.portal_type == 'Comment':
                return ' - '.join([country, sector])
            elif question.portal_type == 'CommentAnswer':
                return ' - '.join([country, 'Coordinator'])

        if userid:
            user = api.user.get(username=userid)
            return user.getProperty('fullname', userid)
        return ''

    def can_add_comment(self):
        return getUtility(IFollowUpPermission)(self.question())

    def can_add_answer(self):
        sm = getSecurityManager()
        question = self.question()
        if question:
            permission = sm.checkPermission('emrt.necd.content: Add CommentAnswer', question)
            questions = [q for q in question.values() if q.portal_type == 'Comment']
            answers = [q for q in question.values() if q.portal_type == 'CommentAnswer']
            return permission and len(questions) > len(answers)
        else:
            return False

    def add_answer_form(self):
        form_instance = AddAnswerForm(self.context, self.request)
        alsoProvides(form_instance, IWrappedForm)
        return form_instance()

    def add_comment_form(self):
        form_instance = AddCommentForm(self.context, self.request)
        alsoProvides(form_instance, IWrappedForm)
        return form_instance()

    def in_conclusions(self):
        state = self.context.get_status()
        return state in [
            'conclusions',
            'conclusions-lr-denied',
            'conclusion-discussion',
            'close-requested',
            'closed',
        ]

    def get_last_editable_thing(self):
        CONCLUSIONS_PHASE_2 = [
            'conclusions',
            'conclusions-lr-denied',
            'conclusion-discussion',
            'close-requested',
        ]

        state = self.context.get_status()
        if state in CONCLUSIONS_PHASE_2:
            return self.context.get_conclusion()
        else:
            question = self.question()
            if question is not None:
                qs = question.get_questions()
                return qs[-1].getObject() if qs else None

        return None

    def update(self):
        context = self.get_last_editable_thing()
        if context is not None:
            if context.can_edit():
                try:
                    history_metadata = self.repo_tool.getHistoryMetadata(context)
                except:
                    history_metadata = None
                if history_metadata:
                    retrieve = history_metadata.retrieve
                    getId = history_metadata.getVersionId
                    history = self.history = []
                    # Count backwards from most recent to least recent
                    for i in xrange(history_metadata.getLength(countPurged=False)-1, -1, -1):
                        version = retrieve(i, countPurged=False)['metadata'].copy()
                        version['version_id'] = getId(i, countPurged=False)
                        history.append(version)
                    dt = getToolByName(self.context, "portal_diff")

                    version1 = self.request.get("one", None)
                    version2 = self.request.get("two", None)

                    if version1 is None and version2 is None:
                        self.history.sort(lambda x,y: cmp(x.get('version_id', ''), y.get('version_id')), reverse=True)
                        version1 = self.history[-1].get('version_id', 'current')
                        if len(self.history) > 1:
                            version2 = self.history[-2].get('version_id', 'current')
                        else:
                            version2 = 'current'
                    elif version1 is None:
                        version1 = 'current'
                    elif version2 is None:
                        version2 = 'current'

                    self.request.set('one', version1)
                    self.request.set('two', version2)
                    changeset = dt.createChangeSet(
                        self.getVersion(version2),
                        self.getVersion(version1),
                        id1=self.versionTitle(version2),
                        id2=self.versionTitle(version1))
                    self.changes = [change for change in changeset.getDiffs()
                                   if not change.same]

    @property
    def repo_tool(self):
        return getToolByName(self.context, "portal_repository")

    def getVersion(self, version):
        context = self.get_last_editable_thing()
        if version == "current":
            return context
        else:
            return self.repo_tool.retrieve(context, int(version)).object

    def versionName(self, version):
        """
        Copied from @@history_view
        Translate the version name. This is needed to allow translation
        when `version` is the string 'current'.
        """
        return _CMFE(version)

    def versionTitle(self, version):
        version_name = self.versionName(version)

        return translate(
            _CMFE(u"version ${version}",
              mapping=dict(version=version_name)),
            context=self.request
        )

    def isChatCurrent(self):
        status = api.content.get_state(self.context)
        if status in ['draft', 'pending']:
            return True
        else:
            return False


class ObservationView(ObservationMixin):

    def get_current_counterparters(self):
        """ Return list of current counterparters,
            if the user can see counterpart action
        """
        actions = [action['action'] for action in self.actions()]
        if not any('counterpart_form' in action for action  in actions):
            return []

        target = self.context
        local_roles = target.get_local_roles()
        users = [
            u[0] for u in local_roles if ROLE_CP in u[1]
        ]
        return [api.user.get(user) for user in users]

    def can_export_observation(self):
        sm = getSecurityManager()
        return sm.checkPermission(
            'emrt.necd.content: Export an Observation', self.context
        )

    def is_projection(self):
        return _is_projection(self.context.aq_parent)


class ExportAsDocView(ObservationMixin):

    def strip_special_chars(self, s):
        """ return s without special chars
        """
        return re.sub('\s+', ' ', s)

    def build_file(self):
        document = Document()

        # Styles
        style = document.styles.add_style('Label Bold', WD_STYLE_TYPE.PARAGRAPH)
        style.font.bold = True
        style = document.styles.add_style('Table Cell', WD_STYLE_TYPE.PARAGRAPH)
        style.font.size = Pt(9)
        style = document.styles.add_style('Table Cell Bold', WD_STYLE_TYPE.PARAGRAPH)
        style.font.size = Pt(9)
        style.font.bold = True

        p = document.add_paragraph('Ref. Number')
        document.add_heading(self.context.getId(), 0)

        p = document.add_paragraph('')
        table = document.add_table(rows=1, cols=5)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Country'
        hdr_cells[0].paragraphs[0].style = "Table Cell Bold"
        hdr_cells[1].text = 'Sector'
        hdr_cells[1].paragraphs[0].style = "Table Cell Bold"
        hdr_cells[2].text = 'Pollutants'
        hdr_cells[2].paragraphs[0].style = "Table Cell Bold"
        hdr_cells[3].text = 'Fuel'
        hdr_cells[3].paragraphs[0].style = "Table Cell Bold"
        hdr_cells[4].text = 'Inventory year'
        hdr_cells[4].paragraphs[0].style = "Table Cell Bold"

        row_cells = table.add_row().cells
        row_cells[0].text = self.context.country_value() or ''
        row_cells[0].paragraphs[0].style = "Table Cell"
        row_cells[1].text = self.context.ghg_source_sectors_value() or ''
        row_cells[1].paragraphs[0].style = "Table Cell"
        row_cells[2].text = self.context.pollutants_value() or ''
        row_cells[2].paragraphs[0].style = "Table Cell"
        row_cells[3].text = get_vocabulary_value(self.context,
            IObservation['fuel'].vocabularyName,
            self.context.fuel
        )
        row_cells[3].paragraphs[0].style = "Table Cell"
        row_cells[4].text = self.context.year or ''
        row_cells[4].paragraphs[0].style = "Table Cell"
        p = document.add_paragraph('')

        document.add_heading('Observation details', level=2)

        p = document.add_paragraph('')
        table = document.add_table(rows=1, cols=4)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Review Year'
        hdr_cells[1].text = 'Parameter'
        hdr_cells[2].text = 'Key category'
        hdr_cells[3].text = 'Last update'
        hdr_cells[0].paragraphs[0].style = "Table Cell Bold"
        hdr_cells[1].paragraphs[0].style = "Table Cell Bold"
        hdr_cells[2].paragraphs[0].style = "Table Cell Bold"
        hdr_cells[3].paragraphs[0].style = "Table Cell Bold"

        row_cells = table.add_row().cells
        row_cells[0].text = "%s" % self.context.review_year or ''
        row_cells[1].text = self.context.parameter_value() or ''
        if self.context.ms_key_category:
            row_cells[2].text = "MS Key category"
        else:
            row_cells[2].text = ""
        row_cells[3].text = self.context.modified().strftime(
            '%d %b %Y, %H:%M CET'
        )
        row_cells[0].paragraphs[0].style = "Table Cell"
        row_cells[1].paragraphs[0].style = "Table Cell"
        row_cells[2].paragraphs[0].style = "Table Cell"
        row_cells[3].paragraphs[0].style = "Table Cell"
        p = document.add_paragraph('')

        p = document.add_paragraph('Description flags', style="Label Bold")
        p = document.add_paragraph(self.context.highlight_value())
        p = document.add_paragraph('Short description by sector expert', style="Label Bold")
        p = document.add_paragraph(self.context.text)

        if self.context.get_status() == 'close-requested':
            document.add_heading('Finish observation', level=2)
            document.add_heading('Observation Finish Requested', level=3)
            p = document.add_paragraph('SE comments on finish observation request:', style="Label Bold")

        conclusion_2 = self.get_conclusion()
        if conclusion_2:
            document.add_page_break()
            document.add_heading('Conclusions', level=2)

            p = document.add_paragraph('Final status of observation:', style="Label Bold")
            p = document.add_paragraph(conclusion_2.reason_value())
            p = document.add_paragraph('Recommendation/internal note:', style="Label Bold")
            p = document.add_paragraph(conclusion_2.text)

        chats = self.get_chat()
        if chats:
            document.add_heading('Q&A', level=2)
            for chat in chats:
                date = chat.effective_date
                sent_info = "Sent on: %s"
                if not date:
                    date = chat.modified()
                    sent_info = "Updated on: %s"

                if chat.portal_type.lower() == 'comment':
                    p = document.add_paragraph(
                        '> %s' % self.strip_special_chars(chat.text)
                    )
                    p = document.add_paragraph(
                        "From TERTs To Member State \t\t %s" % (
                            sent_info % date.strftime('%d %b %Y, %H:%M CET')
                        )
                    )

                if chat.portal_type.lower() == 'commentanswer':
                    p = document.add_paragraph(
                        '< %s' % self.strip_special_chars(chat.text)
                    )
                    p = document.add_paragraph(
                        "From Member State To TERTs \t\t %s" % (
                            sent_info % date.strftime('%d %b %Y, %H:%M CET')
                        )
                    )

        return document

    def render(self):
        """ Export current filters observation in xls
        """
        document = self.build_file()

        response = self.request.response
        response.setHeader(
            "content-type",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
        response.setHeader(
            "Content-disposition", "attachment;filename=" + self.context.getId() + ".docx"
        )

        f = StringIO()
        document.save(f)
        f.seek(0)
        response.setHeader('Content-Length', len(f.getvalue()))
        response.write(f.getvalue())


class AddQuestionForm(Form):

    ignoreContext = True
    fields = field.Fields(IComment).select('text')

    @button.buttonAndHandler(u'Save question')
    def create_question(self, action):
        context = aq_inner(self.context)
        text = self.request.form.get('form.widgets.text', '')
        if not text.strip():
            raise ActionExecutionError(Invalid(u"Question text is empty"))

        qs = self.context.get_values_cat('Question')
        if qs:
            question = qs[0]
        else:
            q_id = context.invokeFactory(
                type_name='Question',
                id='question-1',
                title='Question 1',
            )
            question = context.get(q_id)

        id = str(int(time()))
        item_id = question.invokeFactory(
            type_name='Comment',
            id=id,
        )
        comment = question.get(item_id)
        comment.text = text

        return self.request.response.redirect(context.absolute_url())

    def updateWidgets(self):
        super(AddQuestionForm, self).updateWidgets()
        self.widgets['text'].rows = 15

    def updateActions(self):
        super(AddQuestionForm, self).updateActions()
        for k in self.actions.keys():
            self.actions[k].addClass('standardButton')
            self.actions[k].addClass('defaultWFButton')


class ModificationForm(edit.DefaultEditForm):

    @button.buttonAndHandler(_(u'Save'), name='save')
    def handleApply(self, action):
        if _is_projection(self.context):
            data, errors = self.extractData()

            activity_data = data['activity_data']
            activity_data_type = data['activity_data_type']

            activity_data_validator(self.context, activity_data_type,
                                    activity_data)

        return super(ModificationForm, self).handleApply(self, action)

    @button.buttonAndHandler(_(u'Cancel'), name='cancel')
    def handleCancel(self, action):
        return super(ModificationForm, self).handleCancel(self, action)

    def updateFields(self):
        super(ModificationForm, self).updateFields()

        user = api.user.get_current()
        roles = api.user.get_roles(username=user.getId(), obj=self.context)
        fields = []
        # XXX Needed? Edit rights are controlled by the WF
        if 'Manager' in roles:
            fields = field.Fields(IObservation)
        elif ROLE_SE in roles:
            fields = [f for f in field.Fields(IObservation) if f not in [
                'country',
                'nfr_code',
                'review_year',
                'technical_corrections',
                'closing_comments',
                'closing_deny_comments',

            ]]
        elif ROLE_LR in roles:
            fields = ['text', 'highlight']

        self.fields = field.Fields(IObservation).select(*fields)
        self.groups = [g for g in self.groups if g.label == 'label_schema_default']
        if 'parameter' in fields:
            self.fields['parameter'].widgetFactory = CheckBoxFieldWidget
        if 'highlight' in fields:
            self.fields['highlight'].widgetFactory = CheckBoxFieldWidget
        if 'pollutants' in fields:
            self.fields['pollutants'].widgetFactory = CheckBoxFieldWidget
        if 'activity_data' in fields:
            self.fields['activity_data'].widgetFactory = CheckBoxFieldWidget

    def updateWidgets(self):
        super(ModificationForm, self).updateWidgets()

        w_activity_data = self.widgets['activity_data']

        if _is_projection(self.context):
            self.fields['year'].field.description = YEAR_DESCRIPTION_PROJECTION
            self.widgets['fuel'].mode = interfaces.HIDDEN_MODE
            self.widgets['activity_data_type'].template = \
                Z3ViewPageTemplateFile('templates/widget_activity_type.pt')
            w_activity_data.template = Z3ViewPageTemplateFile(
                'templates/widget_activity.pt'
            )
            w_activity_data.activity_data_registry = json.dumps(
                get_registry_interface_field_data(
                    INECDVocabularies,
                    'activity_data'
                )
            )
        else:
            self.fields['year'].field.description = YEAR_DESCRIPTION_INVENTORY
            self.widgets['activity_data'].mode = interfaces.HIDDEN_MODE
            self.widgets['activity_data_type'].mode = interfaces.HIDDEN_MODE
            self.widgets['pollutants'].template = Z3ViewPageTemplateFile(
                'templates/widget_pollutants.pt'
            )

        self.widgets['highlight'].template = Z3ViewPageTemplateFile(
            'templates/widget_highlight.pt'
        )

    def updateActions(self):
        super(ModificationForm, self).updateActions()
        for k in self.actions.keys():
            self.actions[k].addClass('standardButton')


class AddAnswerForm(Form):

    ignoreContext = True
    fields = field.Fields(ICommentAnswer).select('text')

    @button.buttonAndHandler(u'Save answer')
    def add_answer(self, action):
        text = self.request.form.get('form.widgets.text', '')
        if not text.strip():
            raise ActionExecutionError(Invalid(u"Answer text is empty"))
        observation = aq_inner(self.context)
        questions = [q for q in observation.values() if q.portal_type == 'Question']
        if questions:
            context = questions[0]
        else:
            raise ActionExecutionError(Invalid(u"Invalid context"))
        id = str(int(time()))
        item_id = context.invokeFactory(
            type_name='CommentAnswer',
            id=id,
        )
        comment = context.get(item_id)
        comment.text = text
        action = 'add-answer'
        api.content.transition(obj=context, transition=action)

        return self.request.response.redirect(observation.absolute_url())

    def updateWidgets(self):
        super(AddAnswerForm, self).updateWidgets()
        self.widgets['text'].rows = 15

    def updateActions(self):
        super(AddAnswerForm, self).updateActions()
        for k in self.actions.keys():
            self.actions[k].addClass('standardButton')


class AddAnswerAndRequestComments(BrowserView):
     def render(self):
        observation = aq_inner(self.context)
        questions = [q for q in observation.values() if q.portal_type == 'Question']
        if questions:
            context = questions[0]
        else:
            raise ActionExecutionError(Invalid(u"Invalid context"))

        comments = [q for q in context.values() if q.portal_type == 'Comment']
        answers = [q for q in context.values() if q.portal_type == 'CommentAnswer']

        if (len(comments) == len(answers)):
            status = IStatusMessage(self.request)
            msg = _(u'There is a draft answer created for the question.')
            status.addStatusMessage(msg, "error")
            return self.request.response.redirect(observation.absolute_url())

        context = questions[0]

        text = (
            u'For MS coordinator: please draft, edit and finalize '
            u'here your answer AFTER CLOSING COMMENT within your '
            u'member state expert.'
        )

        id = str(int(time()))
        item_id = context.invokeFactory(
            type_name='CommentAnswer',
            id=id,
        )
        comment = context.get(item_id)
        comment.text = text

        action = 'assign-answerer'
        url = '%s/assign_answerer_form?workflow_action=%s&comment=%s' % (
            context.absolute_url(), action, item_id)

        return self.request.response.redirect(url)


def create_comment(text, question):
    id = str(int(time()))
    item_id = question.invokeFactory(type_name='Comment', id=id)
    comment = question.get(item_id)
    comment.text = text
    return comment


def value_or_error(value, err_text):
    if not value:
        raise ActionExecutionError(Invalid(err_text))
    return value


class AddCommentForm(Form):

    ignoreContext = True
    fields = field.Fields(IComment).select('text')

    @button.buttonAndHandler(u'Add question')
    def create_question(self, action):
        request = self.request
        observation = self.context
        # raising errors before transition as that will
        # cause a transaction.commit
        question = value_or_error(
            observation.get_question(),
            u'Invalid context'
        )
        wid_text = self.widgets['text']

        text = value_or_error(
            wid_text.extract(u'').strip(),
            u'Question text is empty'
        )

        # transition before adding the comment so the transition guard passes
        api.content.transition(obj=question, transition='add-followup-question')
        create_comment(text, question)

        return request.response.redirect(observation.absolute_url())

    def updateWidgets(self):
        super(AddCommentForm, self).updateWidgets()
        self.widgets['text'].rows = 15

    def updateActions(self):
        super(AddCommentForm, self).updateActions()
        for k in self.actions.keys():
            self.actions[k].addClass('standardButton')


class EditConclusionP2AndCloseComments(BrowserView):
    def update(self):
        # Some checks:
        waction = self.request.get('workflow_action')
        if waction != 'finish-comments':
                status = IStatusMessage(self.request)
                msg = u'There was an error, try again please'
                status.addStatusMessage(msg, "error")

    def render(self):
        # Execute the transition
        api.content.transition(
            obj=self.context,
            transition='finish-comments'
        )
        conclusions = self.context.get_values_cat('Conclusions')
        conclusion = conclusions[0]
        url = '%s/edit' % conclusion.absolute_url()
        return self.request.response.redirect(url)


class EditHighlightsForm(edit.DefaultEditForm):
    def updateFields(self):
        super(EditHighlightsForm, self).updateFields()
        self.fields = field.Fields(IObservation).select('highlight')
        self.fields['highlight'].widgetFactory = CheckBoxFieldWidget
        self.groups = [g for g in self.groups if g.label == 'label_schema_default']

    def updateWidgets(self):
        super(EditHighlightsForm, self).updateWidgets()
        self.widgets['highlight'].template = Z3ViewPageTemplateFile(
            'templates/widget_highlight.pt'
        )


class AddConclusions(BrowserView):
    def render(self):
        context = aq_inner(self.context)
        conclusions_folder = self.context.get_values_cat('Conclusions')

        if conclusions_folder:
            conclusions = conclusions_folder[0]
            url = conclusions.absolute_url() + '/edit'

        else:
            url = '{}/++add++Conclusions'.format(context.absolute_url())

        return self.request.response.redirect(url)
